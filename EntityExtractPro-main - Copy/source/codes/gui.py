import tkinter as tk
from tkinter import ttk, scrolledtext, filedialog, messagebox
import threading
from PIL import Image, ImageTk
import fitz  # PyMuPDF
import docx
import os
import logging
from pathlib import Path
import re
from text_analysis import perform_ner_with_logging

class NERApp:
    def __init__(self, root):
        self.root = root
        self.root.title("NER - Natural Entity Recognition")
        self.root.geometry("1200x800")
        self.setup_logging()
        self.setup_gui()
        self.current_page = 0
        self.total_pages = 0
        
    def setup_logging(self):
        Path("logs").mkdir(exist_ok=True)
        logging.basicConfig(filename='logs/ner_app.log', level=logging.INFO,
                            format='%(asctime)s - %(levelname)s - %(message)s')
    
    def setup_gui(self):
        self.root.columnconfigure(0, weight=1)
        self.root.columnconfigure(1, weight=1)
        self.root.rowconfigure(0, weight=1)

        # Left panel
        left_frame = ttk.Frame(self.root)
        left_frame.grid(row=0, column=0, sticky="nsew", padx=10, pady=10)
        left_frame.columnconfigure(0, weight=1)
        left_frame.rowconfigure(1, weight=1)

        self.file_btn = ttk.Button(left_frame, text="Select File", command=self.load_file)
        self.file_btn.grid(row=0, column=0, pady=10)

        self.preview_frame = ttk.Frame(left_frame)
        self.preview_frame.grid(row=1, column=0, sticky="nsew")
        self.preview_frame.columnconfigure(0, weight=1)
        self.preview_frame.rowconfigure(0, weight=1)

        self.canvas = tk.Canvas(self.preview_frame)
        self.canvas.grid(row=0, column=0, sticky="nsew")

        v_scrollbar = ttk.Scrollbar(self.preview_frame, orient="vertical", command=self.canvas.yview)
        v_scrollbar.grid(row=0, column=1, sticky="ns")
        self.canvas.configure(yscrollcommand=v_scrollbar.set)

        self.content_frame = ttk.Frame(self.canvas)
        self.canvas.create_window((0, 0), window=self.content_frame, anchor="nw")

        self.content_frame.bind("<Configure>", lambda e: self.canvas.configure(scrollregion=self.canvas.bbox("all")))

        self.run_btn = ttk.Button(left_frame, text="Run Analysis", command=self.run_analysis)
        self.run_btn.grid(row=2, column=0, pady=10)

        # Page navigation (for PDFs and multi-page images)
        self.nav_frame = ttk.Frame(left_frame)
        self.nav_frame.grid(row=3, column=0, pady=10)
        self.prev_btn = ttk.Button(self.nav_frame, text="Previous", command=self.prev_page)
        self.prev_btn.grid(row=0, column=0, padx=5)
        self.next_btn = ttk.Button(self.nav_frame, text="Next", command=self.next_page)
        self.next_btn.grid(row=0, column=1, padx=5)
        self.page_label = ttk.Label(self.nav_frame, text="Page: 0 / 0")
        self.page_label.grid(row=0, column=2, padx=5)

        # Right panel
        right_frame = ttk.Frame(self.root)
        right_frame.grid(row=0, column=1, sticky="nsew", padx=10, pady=10)
        right_frame.columnconfigure(0, weight=1)
        right_frame.rowconfigure(0, weight=1)
        right_frame.rowconfigure(1, weight=1)

        self.results_text = scrolledtext.ScrolledText(right_frame, wrap=tk.WORD, width=40, height=15)
        self.results_text.grid(row=0, column=0, sticky="nsew")

        self.summary_text = scrolledtext.ScrolledText(right_frame, wrap=tk.WORD, width=40, height=15)
        self.summary_text.grid(row=1, column=0, sticky="nsew")

        self.save_btn = ttk.Button(right_frame, text="Save Results", command=self.save_results)
        self.save_btn.grid(row=2, column=0, pady=10)

        # Status bar
        self.status_var = tk.StringVar()
        self.status_bar = ttk.Label(self.root, textvariable=self.status_var, relief=tk.SUNKEN, anchor=tk.W)
        self.status_bar.grid(row=1, column=0, columnspan=2, sticky="ew")

    def load_file(self):
        try:
            self.file_path = filedialog.askopenfilename(filetypes=[
                ("All supported files", "*.pdf;*.docx;*.txt;*.png;*.jpg;*.jpeg"),
                ("PDF files", "*.pdf"),
                ("Word files", "*.docx"),
                ("Text files", "*.txt"),
                ("Image files", "*.png;*.jpg;*.jpeg")
            ])
            if not self.file_path:
                return
            
            self.status_var.set("Loading file...")
            self.root.update_idletasks()
            
            self.clear_preview()
            
            if self.file_path.endswith('.pdf'):
                self.load_pdf()
            elif self.file_path.endswith('.docx'):
                self.load_docx()
            elif self.file_path.endswith('.txt'):
                self.load_txt()
            elif self.file_path.lower().endswith(('.png', '.jpg', '.jpeg')):
                self.load_image()
            else:
                raise ValueError("Unsupported file type")
            
            self.status_var.set("File loaded successfully.")
            logging.info(f"Loaded file: {self.file_path}")
        except Exception as e:
            messagebox.showerror("Error", f"An error occurred while loading the file: {str(e)}")
            logging.error(f"Error loading file: {str(e)}")
            self.status_var.set("Error loading file.")

    def clear_preview(self):
        for widget in self.content_frame.winfo_children():
            widget.destroy()

    def load_pdf(self):
        self.doc = fitz.open(self.file_path)
        self.total_pages = len(self.doc)
        self.current_page = 0
        self.display_pdf_page()

    def display_pdf_page(self):
        self.clear_preview()
        page = self.doc.load_page(self.current_page)
        pix = page.get_pixmap()
        img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
        img_tk = ImageTk.PhotoImage(image=img)
        label = ttk.Label(self.content_frame, image=img_tk)
        label.image = img_tk
        label.pack()
        self.update_page_label()

    def load_docx(self):
        doc = docx.Document(self.file_path)
        text = "\n".join([para.text for para in doc.paragraphs])
        self.display_text(text)

    def load_txt(self):
        with open(self.file_path, 'r', encoding='utf-8') as file:
            text = file.read()
        self.display_text(text)

    def load_image(self):
        img = Image.open(self.file_path)
        img_tk = ImageTk.PhotoImage(image=img)
        label = ttk.Label(self.content_frame, image=img_tk)
        label.image = img_tk
        label.pack()

    def display_text(self, text):
        self.text_widget = tk.Text(self.content_frame, wrap=tk.WORD)
        self.text_widget.pack(fill=tk.BOTH, expand=True)
        self.text_widget.insert(tk.END, text)
        self.text_widget.config(state=tk.DISABLED)

    def update_page_label(self):
        self.page_label.config(text=f"Page: {self.current_page + 1} / {self.total_pages}")

    def prev_page(self):
        if self.current_page > 0:
            self.current_page -= 1
            self.display_pdf_page()

    def next_page(self):
        if self.current_page < self.total_pages - 1:
            self.current_page += 1
            self.display_pdf_page()

    def run_analysis(self):
        if not hasattr(self, 'file_path'):
            messagebox.showwarning("No File", "Please load a file before running analysis.")
            return
        
        self.status_var.set("Running analysis...")
        self.root.update_idletasks()
        
        threading.Thread(target=self._run_analysis_thread, daemon=True).start()

    def _run_analysis_thread(self):
        try:
            if self.file_path.endswith('.pdf'):
                text = "\n".join([page.get_text() for page in self.doc])
            elif self.file_path.endswith('.docx'):
                doc = docx.Document(self.file_path)
                text = "\n".join([para.text for para in doc.paragraphs])
            elif self.file_path.endswith('.txt'):
                with open(self.file_path, 'r', encoding='utf-8') as file:
                    text = file.read()
            else:  # Image file
                text = "Image analysis not implemented"  # You might want to use OCR here

            entities, categories, summary = perform_ner_with_logging(text)
            
            self.root.after(0, self._update_results, categories, summary)
            self.root.after(0, self.highlight_entities, entities)
            logging.info("Analysis completed successfully")
        except Exception as e:
            self.root.after(0, messagebox.showerror, "Error", f"An error occurred during analysis: {str(e)}")
            logging.error(f"Error during analysis: {str(e)}")
        finally:
            self.root.after(0, self.status_var.set, "Analysis complete.")

    def _update_results(self, categories, summary):
        self.results_text.delete(1.0, tk.END)
        for category, items in categories.items():
            self.results_text.insert(tk.END, f"{category}:\n")
            for item in set(items):  # Use set to remove duplicates
                self.results_text.insert(tk.END, f"  - {item}\n")
            self.results_text.insert(tk.END, "\n")
        
        self.summary_text.delete(1.0, tk.END)
        self.summary_text.insert(tk.END, summary)

    def highlight_entities(self, entities):
        if hasattr(self, 'text_widget'):
            self.text_widget.config(state=tk.NORMAL)
            for text, label, start, end in entities:
                self.text_widget.tag_add(label, f"1.0+{start}c", f"1.0+{end}c")
                self.text_widget.tag_config(label, background=self.get_color_for_label(label))
            self.text_widget.config(state=tk.DISABLED)
        elif self.file_path.endswith('.pdf'):
            self.highlight_pdf_entities(entities)

    def highlight_pdf_entities(self, entities):
        Path("highlighted").mkdir(exist_ok=True)
        output_file = Path("highlighted") / (Path(self.file_path).stem + "_highlighted.pdf")
        
        doc = fitz.open(self.file_path)
        for page in doc:
            for text, label, _, _ in entities:
                areas = page.search_for(text)
                for rect in areas:
                    highlight = page.add_highlight_annot(rect)
                    highlight.set_colors(stroke=self.get_color_for_label(label, format="rgb"))
                    highlight.update()
        
        doc.save(str(output_file))
        messagebox.showinfo("Highlighting Complete", f"Highlighted PDF saved as: {output_file}")

    def get_color_for_label(self, label, format="hex"):
        color_map = {
            "PERSON": ("#FFA07A", (1, 0.627, 0.478)),    # Light Salmon
            "ORG": ("#98FB98", (0.596, 0.984, 0.596)),   # Pale Green
            "GPE": ("#87CEFA", (0.529, 0.808, 0.980)),   # Light Sky Blue
            "DATE": ("#DDA0DD", (0.867, 0.627, 0.867)),  # Plum
            "CARDINAL": ("#F0E68C", (0.941, 0.902, 0.549)), # Khaki
            "NORP": ("#20B2AA", (0.125, 0.698, 0.667)),  # Light Sea Green
            "MONEY": ("#90EE90", (0.565, 0.933, 0.565)), # Light Green
            "RANK": ("#FFB6C1", (1, 0.714, 0.757)),      # Light Pink
            "ID": ("#E6E6FA", (0.902, 0.902, 0.980)),    # Lavender
            "PHONE": ("#FFDAB9", (1, 0.855, 0.725)),     # Peach Puff
            "EVENT": ("#B0E0E6", (0.690, 0.878, 0.902))  # Powder Blue
        }
        default_color = ("#DCDCDC", (0.863, 0.863, 0.863)) # Gainsboro (light gray)
        
        if format == "hex":
            return color_map.get(label, default_color)[0]
        else:  # rgb
            return color_map.get(label, default_color)[1]

    def save_results(self):
        if not self.results_text.get(1.0, tk.END).strip():
            messagebox.showwarning("No Results", "Please run analysis before saving results.")
            return
        
        file_path = filedialog.asksaveasfilename(defaultextension=".txt",
                                                 filetypes=[("Text files", "*.txt")])
        if file_path:
            try:
                with open(file_path, 'w') as f:
                    f.write("Entity Recognition Results:\n\n")
                    f.write(self.results_text.get(1.0, tk.END))
                    f.write("\nAI-Generated Summary:\n\n")
                    f.write(self.summary_text.get(1.0, tk.END))
                messagebox.showinfo("Success", "Results saved successfully.")
                logging.info(f"Results saved to: {file_path}")
            except Exception as e:
                messagebox.showerror("Error", f"An error occurred while saving results: {str(e)}")
                logging.error(f"Error saving results: {str(e)}")

if __name__ == "__main__":
    root = tk.Tk()
    app = NERApp(root)
    root.mainloop()